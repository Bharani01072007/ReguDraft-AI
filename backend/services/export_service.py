import io
from docx import Document as DocxDocument
from fpdf import FPDF

class ExportService:
    def export_to_docx(self, markdown_content: str) -> bytes:
        doc = DocxDocument()
        lines = markdown_content.split("\n")
        
        for line in lines:
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

        file_stream = io.BytesIO()
        doc.save(file_stream)
        return file_stream.getvalue()

    def _clean_unicode_for_latin1(self, text: str) -> str:
        replacements = {
            "\u2018": "'",
            "\u2019": "'",
            "\u201c": '"',
            "\u201d": '"',
            "\u2013": "-",
            "\u2014": "-",
            "\u2265": ">=",
            "\u2264": "<=",
            "\u2022": "*",
            "\u00a0": " ",
            "\u00b5": "u",
            "\u00ae": "(R)",
            "\u00a9": "(C)",
            "\u2212": "-",
        }
        for orig, rep in replacements.items():
            text = text.replace(orig, rep)
        # Force encode to latin-1, replacing any other unencodable characters with '?'
        return text.encode("latin-1", "replace").decode("latin-1")

    def export_to_pdf(self, markdown_content: str) -> bytes:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Helvetica", size=12)

        cleaned_content = self._clean_unicode_for_latin1(markdown_content)
        lines = cleaned_content.split("\n")
        for line in lines:
            # Basic styling based on markdown
            if line.startswith("# "):
                pdf.set_font("Helvetica", style="B", size=16)
                pdf.cell(pdf.epw, 10, text=line[2:])
                pdf.ln(12)
            elif line.startswith("## "):
                pdf.set_font("Helvetica", style="B", size=14)
                pdf.cell(pdf.epw, 10, text=line[3:])
                pdf.ln(10)
            elif line.startswith("### "):
                pdf.set_font("Helvetica", style="B", size=12)
                pdf.cell(pdf.epw, 10, text=line[4:])
                pdf.ln(8)
            elif line.strip() == "":
                pdf.ln(4)
            else:
                pdf.set_font("Helvetica", size=10)
                pdf.multi_cell(pdf.epw, 6, text=line)
                pdf.ln(2)
                
        return pdf.output()

    def export_to_txt(self, markdown_content: str) -> bytes:
        return markdown_content.encode("utf-8")

export_service = ExportService()
